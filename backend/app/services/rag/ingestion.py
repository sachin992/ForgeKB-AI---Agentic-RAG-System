import hashlib
import json
from collections.abc import Callable
from pathlib import Path

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from app.core.config import settings
from app.db.models import DocumentRegistry
from app.db.session import SessionLocal


def _vector_ids_for_sources(store: FAISS, sources: set[str]) -> list[str]:
    """Collect vector docstore ids that belong to the given sources."""
    if not sources:
        return []
    docstore = getattr(store, "docstore", None)
    doc_map = getattr(docstore, "_dict", {}) if docstore is not None else {}
    ids: list[str] = []
    for doc_id, doc in doc_map.items():
        metadata = getattr(doc, "metadata", {}) or {}
        if metadata.get("source", "") in sources:
            ids.append(str(doc_id))
    return ids


def _loader_for_file(path: Path):
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return TextLoader(str(path), encoding="utf-8")
    if suffix == ".pdf":
        return PyPDFLoader(str(path))
    return None


def _load_docx(path: Path) -> list[Document]:
    doc = DocxDocument(str(path))
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text, metadata={"source": str(path)})]


def _read_manifest() -> dict:
    manifest_path = Path(settings.metadata_store_path)
    if not manifest_path.exists():
        return {"documents": [], "chunks": []}
    return json.loads(manifest_path.read_text(encoding="utf-8"))


def _write_manifest(payload: dict) -> None:
    manifest_path = Path(settings.metadata_store_path)
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")


def _file_hash(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        while True:
            data = f.read(8192)
            if not data:
                break
            h.update(data)
    return h.hexdigest()


def _to_int(value: object | None, default: int = 0) -> int:
    try:
        return int(value) if value is not None else default
    except Exception:
        return default


def _doc_meta_from_row(row: DocumentRegistry, file_hash: str) -> dict:
    return {
        "source": row.file_path,
        "display_name": row.display_name or Path(row.file_path).name,
        "recency_id": row.id,
        "hash": file_hash,
        "version": row.version,
        "is_deleted": False,
        "owner_user_id": row.owner_user_id,
        "visibility": row.visibility or "private",
    }


def _chunk_meta_from_doc(doc_meta: dict, idx: int) -> dict:
    source = doc_meta["source"]
    return {
        "source": source,
        "chunk_id": f"{Path(source).name}:{idx}",
        "parent_id": source,
        "hash": doc_meta["hash"],
        "display_name": doc_meta.get("display_name", Path(source).name),
        "recency_id": doc_meta.get("recency_id", 0),
        "owner_user_id": doc_meta.get("owner_user_id"),
        "visibility": doc_meta.get("visibility", "private"),
    }


def run_ingestion_pipeline(
    progress_callback: Callable[[int, str, dict | None], None] | None = None,
    affected_sources: set[str] | None = None,
) -> dict:
    uploads_dir = Path(settings.uploads_dir)
    uploads_dir.mkdir(parents=True, exist_ok=True)

    splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=120)
    embeddings = HuggingFaceEmbeddings(model_name=settings.embedding_model)

    db: Session = SessionLocal()
    try:
        if progress_callback:
            progress_callback(5, "scanning", {"note": "Scanning allowed sources"})

        active_rows = db.query(DocumentRegistry).filter(DocumentRegistry.is_deleted == False).all()
        active_by_source = {row.file_path: row for row in active_rows}

        old_manifest = _read_manifest()
        old_docs = {d.get("source", ""): d for d in old_manifest.get("documents", [])}
        old_chunks = old_manifest.get("chunks", [])
        old_chunk_sources = {c.get("source", "") for c in old_chunks}

        if affected_sources is None:
            target_sources = set(active_by_source.keys()) | {
                src for src in old_docs.keys() if src and src not in active_by_source
            }
        else:
            target_sources = {src for src in affected_sources if src}

        files = [Path(src) for src in sorted(target_sources) if src in active_by_source and Path(src).exists()]
        total_files = len(files)
        if progress_callback:
            progress_callback(10, "loading", {"total_files": total_files, "target_sources": len(target_sources)})

        unchanged_sources: set[str] = set()
        new_sources: set[str] = set()
        modified_sources: set[str] = set()

        processed_files = 0
        for path in files:
            source = str(path)
            row = active_by_source.get(source)
            if not row:
                continue
            file_hash = _file_hash(path)
            old_doc = old_docs.get(source)

            if old_doc and old_doc.get("hash") == file_hash and source in old_chunk_sources:
                unchanged_sources.add(source)
            elif old_doc:
                modified_sources.add(source)
            else:
                new_sources.add(source)

            row.file_hash = file_hash
            db.add(row)
            processed_files += 1
            if progress_callback and total_files > 0:
                pct = 10 + int((processed_files / total_files) * 50)
                progress_callback(min(pct, 60), "chunking", {"file": path.name, "processed": processed_files})

        deleted_sources = {src for src in target_sources if src in old_docs and src not in active_by_source}

        all_chunks: list[Document] = []
        documents: list[dict] = []

        def _load_split_docs(path: Path, row: DocumentRegistry, file_hash: str) -> tuple[list[Document], dict]:
            if path.suffix.lower() == ".docx":
                docs = _load_docx(path)
            else:
                loader = _loader_for_file(path)
                if loader is None:
                    return [], _doc_meta_from_row(row, file_hash)
                docs = loader.load()
            split_docs = splitter.split_documents(docs)
            doc_meta = _doc_meta_from_row(row, file_hash)
            for idx, doc in enumerate(split_docs):
                doc.metadata.update(_chunk_meta_from_doc(doc_meta, idx))
            return split_docs, doc_meta

        faiss_path = Path(settings.faiss_dir)
        has_index = faiss_path.exists() and (faiss_path / "index.faiss").exists()

        # Fast path: no data change for targeted sources, keep existing vectors as-is.
        if not new_sources and not modified_sources and not deleted_sources and has_index:
            for source, doc_meta in old_docs.items():
                if source in active_by_source:
                    documents.append(doc_meta)
            db.commit()
            if progress_callback:
                progress_callback(95, "dedupe_skip", {"skipped_sources": len(unchanged_sources)})
            return {
                "document_count": len(documents),
                "chunk_count": len(old_chunks),
                "vector_index_ready": True,
                "skipped_existing_embeddings": len(unchanged_sources),
            }

        # Targeted incremental path: only affected sources are changed in FAISS + manifest.
        if has_index:
            if progress_callback:
                progress_callback(
                    75,
                    "embedding_incremental",
                    {
                        "new_files": len(new_sources),
                        "modified_files": len(modified_sources),
                        "deleted_files": len(deleted_sources),
                    },
                )

            store = FAISS.load_local(settings.faiss_dir, embeddings, allow_dangerous_deserialization=True)
            new_chunks: list[Document] = []
            affected_for_reindex = set(new_sources) | set(modified_sources)
            affected_for_removal = affected_for_reindex | set(deleted_sources)

            vector_ids = _vector_ids_for_sources(store, affected_for_removal)
            if vector_ids:
                store.delete(vector_ids)

            for source in sorted(affected_for_reindex):
                row = active_by_source[source]
                split_docs, doc_meta = _load_split_docs(Path(source), row, row.file_hash)
                documents.append(doc_meta)
                new_chunks.extend(split_docs)

            for source, doc_meta in old_docs.items():
                if source in affected_for_removal:
                    continue
                if source in active_by_source:
                    documents.append(doc_meta)

            if new_chunks:
                store.add_documents(new_chunks)
            store.save_local(settings.faiss_dir)

            manifest_chunks = [c for c in old_chunks if c.get("source", "") not in affected_for_removal]
            manifest_chunks.extend(
                [
                    {
                        "chunk_id": c.metadata.get("chunk_id", ""),
                        "source": c.metadata.get("source", ""),
                        "parent_id": c.metadata.get("parent_id", ""),
                        "hash": c.metadata.get("hash", ""),
                        "display_name": c.metadata.get("display_name", Path(c.metadata.get("source", "")).name),
                        "recency_id": _to_int(c.metadata.get("recency_id"), 0),
                        "owner_user_id": _to_int(c.metadata.get("owner_user_id"), 0),
                        "visibility": c.metadata.get("visibility", "private"),
                        "text": c.page_content,
                    }
                    for c in new_chunks
                ]
            )
            _write_manifest({"documents": documents, "chunks": manifest_chunks})

            db.commit()
            if progress_callback:
                progress_callback(
                    98,
                    "finalizing",
                    {
                        "incremental": True,
                        "new_chunks": len(new_chunks),
                        "removed_sources": len(affected_for_removal),
                    },
                )
            return {
                "document_count": len(documents),
                "chunk_count": len(manifest_chunks),
                "vector_index_ready": True,
                "skipped_existing_embeddings": len(unchanged_sources),
            }

        # Full rebuild path: cold start when index is missing.
        for row in active_rows:
            path = Path(row.file_path)
            if not path.exists() or not row.file_hash:
                continue
            split_docs, doc_meta = _load_split_docs(path, row, row.file_hash)
            documents.append(doc_meta)
            all_chunks.extend(split_docs)

        if all_chunks:
            if progress_callback:
                progress_callback(75, "embedding", {"chunk_count": len(all_chunks)})
            vectorstore = FAISS.from_documents(all_chunks, embeddings)
            vectorstore.save_local(settings.faiss_dir)

        if progress_callback:
            progress_callback(90, "manifest", {"note": "Writing chunk manifest"})

        manifest_chunks = [
            {
                "chunk_id": c.metadata.get("chunk_id", ""),
                "source": c.metadata.get("source", ""),
                "parent_id": c.metadata.get("parent_id", ""),
                "hash": c.metadata.get("hash", ""),
                "display_name": c.metadata.get("display_name", Path(c.metadata.get("source", "")).name),
                "recency_id": _to_int(c.metadata.get("recency_id"), 0),
                "owner_user_id": _to_int(c.metadata.get("owner_user_id"), 0),
                "visibility": c.metadata.get("visibility", "private"),
                "text": c.page_content,
            }
            for c in all_chunks
        ]
        _write_manifest({"documents": documents, "chunks": manifest_chunks})

        db.commit()

        if progress_callback:
            progress_callback(98, "finalizing", {"document_count": len(documents), "chunk_count": len(manifest_chunks)})

        return {
            "document_count": len(documents),
            "chunk_count": len(manifest_chunks),
            "vector_index_ready": bool(all_chunks),
            "skipped_existing_embeddings": len(unchanged_sources),
        }
    finally:
        db.close()
